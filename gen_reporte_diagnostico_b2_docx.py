#!/usr/bin/env python3
"""Genera REPORTE_DIAGNOSTICO_COHORTE_B2.docx — editable en Word.
Para que la profesora nueva llene 1 página por estudiante + summary del grupo
después de la sesión diagnóstica de 2h."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

D = os.path.dirname(os.path.abspath(__file__))
EVAL_DIR = os.path.join(D, 'recursos', 'evaluaciones')
DOCX_PATH = os.path.join(EVAL_DIR, 'REPORTE_DIAGNOSTICO_COHORTE_B2.docx')

doc = Document()

# Estilo Normal
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return p


def add_intro_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.size = Pt(10)


def add_field_table(rows, col_widths=(5, 12)):
    """rows = [(label, default_value)]"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    t.columns[0].width = Cm(col_widths[0])
    t.columns[1].width = Cm(col_widths[1])
    for i, (label, val) in enumerate(rows):
        cell_label = t.rows[i].cells[0]
        cell_val = t.rows[i].cells[1]
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        cell_val.paragraphs[0].text = val
    return t


def add_writing_block(label, lines=2):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    for _ in range(lines):
        line_p = doc.add_paragraph('_' * 100)
        line_p.paragraph_format.space_after = Pt(2)


def add_rating_block(label, dims, scale='1 = lowest, 5 = highest'):
    """label + tabla con dimensiones para puntuar."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    note = doc.add_paragraph()
    note_run = note.add_run(f'Scale: {scale}')
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    t = doc.add_table(rows=len(dims) + 1, cols=2)
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    t.columns[0].width = Cm(8)
    t.columns[1].width = Cm(9)
    # Header
    hdr_label = t.rows[0].cells[0].paragraphs[0]
    r = hdr_label.add_run('Dimension')
    r.bold = True
    hdr_score = t.rows[0].cells[1].paragraphs[0]
    r = hdr_score.add_run('Score (1-5) + brief comment')
    r.bold = True
    for i, dim in enumerate(dims, 1):
        t.rows[i].cells[0].paragraphs[0].text = dim
        t.rows[i].cells[1].paragraphs[0].text = ''
    return t


def add_separator():
    p = doc.add_paragraph('─' * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)


# =====================================================================
# TÍTULO
# =====================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('B2 COHORT DIAGNOSTIC REPORT')
title_run.bold = True
title_run.font.size = Pt(18)

add_intro_quote(
    'Submitted by the new B2 teacher after the 2-hour diagnostic session. '
    'Goal: rebuild documentation of the cohort that the previous teacher did not maintain. '
    'This report is the foundation for the next 6 months of B2 lessons. '
    'Submit to coordination same night as the session.'
)

add_separator()

# =====================================================================
# SESSION OVERVIEW
# =====================================================================

add_heading('SESSION OVERVIEW', level=2)
add_field_table([
    ('Teacher (you)', ''),
    ('Date of session', ''),
    ('Time start / end', ''),
    ('Modality', '☐ In-person   ☐ Virtual'),
    ('Number of students present', '___ / 5'),
    ('Absent students (names)', ''),
])

add_separator()

# =====================================================================
# STUDENT 1-5
# =====================================================================

for n in range(1, 6):
    add_heading(f'STUDENT {n}', level=2)

    add_field_table([
        ('Full name', ''),
        ('Real estimated level', '☐ B2 solid   ☐ B2 borderline   ☐ B1+   ☐ Below B1'),
        ('Should they stay in B2?', '☐ Yes   ☐ Yes with extra support   ☐ Recommend re-evaluation'),
    ])

    add_paragraph('')

    # Strengths
    add_writing_block('3 STRENGTHS observed (specific, with examples):', lines=4)
    add_intro_quote(
        'Example: "Used connectors like however and on the other hand fluently in debate." '
        'NOT generic ("good vocabulary"). Specific.'
    )

    add_paragraph('')

    # Weaknesses
    add_writing_block('3 WEAKNESSES / areas of opportunity (specific):', lines=4)
    add_intro_quote(
        'Example: "Confuses Past Simple with Present Perfect when narrating personal stories." '
        'NOT vague ("needs grammar work").'
    )

    add_paragraph('')

    # Ratings
    add_rating_block(
        'Skill ratings:',
        [
            'Pronunciation',
            'Fluency (smooth flow, no major pauses)',
            'Vocabulary range',
            'Grammar accuracy',
            'Argumentation / structure',
            'Confidence under pressure',
            'Listening (responds well to others)',
        ]
    )

    add_paragraph('')

    # Quotes captured
    add_writing_block('Best moment / quote that surprised you (in their words):', lines=2)
    add_writing_block('Worst struggle / quote where they broke down:', lines=2)

    add_paragraph('')

    # Emotional read
    p = doc.add_paragraph()
    p.add_run('Emotional read: ').bold = True
    options = doc.add_paragraph()
    options.paragraph_format.left_indent = Cm(0.5)
    options.add_run('☐  Motivated and engaged    ☐  Confident but coasting    '
                    '☐  Insecure / anxious    ☐  Disengaged / checked out    ☐  Defensive')

    add_paragraph('')

    # Recommendation
    add_writing_block('SPECIFIC recommendation: what should this student work on FIRST?', lines=3)

    add_paragraph('')

    # Personal context
    add_writing_block('Personal/professional context they shared (motivation, profession, goals):', lines=3)

    add_paragraph('')

    # Free notes
    add_writing_block('Free notes / red flags / anything else:', lines=3)

    add_separator()

# =====================================================================
# GROUP SUMMARY
# =====================================================================

add_heading('GROUP SUMMARY', level=1)

add_writing_block('Group cohesion (do they support each other? competitive? distant?):', lines=3)
add_writing_block('Group dynamics (who leads? who is silent? any tensions?):', lines=3)
add_writing_block('Overall level of the cohort (in your honest opinion):', lines=2)

add_paragraph('')

# Top 3 priorities
add_heading('TOP 3 PEDAGOGICAL PRIORITIES FOR THE NEXT 5 CLASSES', level=2)
add_intro_quote(
    'What does THIS group need most? Not generic — specific to what you observed today.'
)

add_writing_block('Priority 1:', lines=3)
add_writing_block('Priority 2:', lines=3)
add_writing_block('Priority 3:', lines=3)

add_paragraph('')

# Alerts
add_heading('ALERTS FOR COORDINATION', level=2)
add_intro_quote(
    'Anything coordination should know urgently. A student at risk of dropping. '
    'A specific situation that needs intervention. A request you have.'
)

add_writing_block('Alert / Request:', lines=4)

add_separator()

# =====================================================================
# CIERRE
# =====================================================================

add_heading('SUBMISSION', level=2)

add_field_table([
    ('Time finished writing this report', ''),
    ('Sent to coordination at', '☐ WhatsApp   ☐ Email   ☐ Both     Time: ___'),
    ('Teacher signature (digital ok)', ''),
])

add_paragraph('')

p = doc.add_paragraph()
run = p.add_run(
    'Submit this report to coordination the same night as the diagnostic session. '
    'The next class plan depends on this data.'
)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Footer
add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'PROPIEDAD INTELECTUAL DE HEIIU — USO INTERNO. '
    'No compartir fuera del equipo de coordinación y profesores.'
)
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(DOCX_PATH)
print(f'OK: {DOCX_PATH}')
