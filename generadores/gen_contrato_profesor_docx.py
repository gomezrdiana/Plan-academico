#!/usr/bin/env python3
import os as _hos, sys as _hsys
_hsys.path.insert(0, _hos.path.dirname(_hos.path.dirname(_hos.path.abspath(__file__))))
"""Genera DOCX (Word) de la plantilla v2 del contrato del profesor.
Uso: python gen_contrato_profesor_docx.py
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

D = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONTRATOS_DIR = os.path.join(D, 'contratos')
MD_PATH = os.path.join(CONTRATOS_DIR, 'PLANTILLA_CONTRATO_HORA_CATEDRA_HEIIU_2026_v2.md')
DOCX_PATH = os.path.join(CONTRATOS_DIR, 'PLANTILLA_CONTRATO_HORA_CATEDRA_HEIIU_2026_v2.docx')


def add_run_with_inline_formatting(paragraph, text):
    """Procesa **bold**, *italic* y los aplica al runs."""
    # Pattern combinado: **texto** o *texto*
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()

    # Configurar estilo Normal por defecto
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = text.split('\n')
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        # Línea vacía -> salto
        if not line.strip():
            i += 1
            continue

        # Separador horizontal
        if line.strip() == '---':
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading 1
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            quote_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Lista numerada (1. ... 18.)
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_run_with_inline_formatting(p, m.group(2))
            i += 1
            continue

        # Lista con bullet (a) b) etc., o "- ")
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_run_with_inline_formatting(p, line[2:])
            i += 1
            continue

        # Lista a) b) c) ...
        m = re.match(r'^([a-z])\)\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_run_with_inline_formatting(p, f'{m.group(1)}) {m.group(2)}')
            i += 1
            continue

        # Párrafo normal con formato inline
        p = doc.add_paragraph()
        # Justificar
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run_with_inline_formatting(p, line)
        i += 1

    doc.save(docx_path)
    print(f'OK: {docx_path}')


if __name__ == '__main__':
    md_to_docx(MD_PATH, DOCX_PATH)
