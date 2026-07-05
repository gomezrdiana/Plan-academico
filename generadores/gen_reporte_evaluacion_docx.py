#!/usr/bin/env python3
import os as _hos, sys as _hsys
_hsys.path.insert(0, _hos.path.dirname(_hos.path.dirname(_hos.path.abspath(__file__))))
"""Genera REPORTE_EVALUACION_PROSPECTO.docx — formato editable para que la evaluadora llene en Word."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EVAL_DIR = os.path.join(D, 'recursos', 'evaluaciones')
DOCX_PATH = os.path.join(EVAL_DIR, 'REPORTE_EVALUACION_PROSPECTO.docx')

doc = Document()

# Configurar estilo Normal
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# ===== Helpers =====

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_intro_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.size = Pt(10)


def add_field_table(rows):
    """rows = [(label, default_value_or_empty)]"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    t.columns[0].width = Cm(6)
    t.columns[1].width = Cm(11)
    for i, (label, val) in enumerate(rows):
        cell_label = t.rows[i].cells[0]
        cell_val = t.rows[i].cells[1]
        # Label en bold
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        # Value (vacío para llenar)
        cell_val.paragraphs[0].text = val
    return t


def add_checkbox_options(label, options, multiple=False):
    """Bloque con etiqueta + opciones marcables (☐)."""
    p = doc.add_paragraph()
    run_lbl = p.add_run(label)
    run_lbl.bold = True
    note = '   (marcar UNA con X o ☑)' if not multiple else '   (marcar las que apliquen)'
    run_note = p.add_run(note)
    run_note.italic = True
    run_note.font.size = Pt(9)
    run_note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for opt in options:
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Cm(0.7)
        op.add_run('☐  ').font.size = Pt(12)
        op.add_run(opt)


def add_writing_block(label, lines=2, hint=''):
    """Bloque con etiqueta + N líneas para escribir."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    if hint:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        rh = p2.add_run(hint)
        rh.italic = True
        rh.font.size = Pt(9)
        rh.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for _ in range(lines):
        line_p = doc.add_paragraph('_' * 100)
        line_p.paragraph_format.space_after = Pt(2)


def add_quote_block(label, lines=2):
    """Bloque para 'frase textual del prospecto'."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    qp.paragraph_format.right_indent = Cm(0.5)
    run = qp.add_run('"')
    run.font.size = Pt(14)
    run.bold = True
    for _ in range(lines):
        line_p = doc.add_paragraph()
        line_p.paragraph_format.left_indent = Cm(1)
        line_p.add_run('_' * 90)
    qp_close = doc.add_paragraph()
    qp_close.paragraph_format.left_indent = Cm(1)
    run = qp_close.add_run('"')
    run.font.size = Pt(14)
    run.bold = True


def add_separator():
    p = doc.add_paragraph('─' * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)


# ===== TÍTULO =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('REPORTE DE EVALUACIÓN — PROSPECTO HEIIU')
title_run.bold = True
title_run.font.size = Pt(18)

add_intro_quote(
    'Para llenar por la evaluadora INMEDIATAMENTE al colgar/cerrar la evaluación. '
    'Tiempo de llenado: 3-5 min máximo. '
    'Entregable obligatorio para coordinación + asesor comercial. '
    'Sin este reporte, el asesor improvisa y el cliente se enfría.'
)

add_separator()

# ===== SECCIÓN 1 =====
add_heading('1. DATOS BÁSICOS', level=2)
add_field_table([
    ('Nombre completo del prospecto', ''),
    ('Edad aproximada', ''),
    ('Profesión / ocupación', ''),
    ('Fecha de la evaluación', ''),
    ('Hora inicio / fin', ''),
    ('Modalidad', '☐ Presencial   ☐ Virtual (Meet)'),
    ('Evaluadora', ''),
    ('Asesor comercial responsable', ''),
])

add_separator()

# ===== SECCIÓN 2 =====
add_heading('2. DIAGNÓSTICO DE NIVEL', level=2)

add_checkbox_options('Nivel diagnosticado:', [
    'A1', 'A2', 'B1', 'B2'
])

add_paragraph('')  # espacio

add_checkbox_options('Subnivel / certeza del diagnóstico:', [
    '100% claro',
    '80% — está al borde de subir/bajar',
    '60% — necesita 2do test en 2 semanas'
])

add_checkbox_options('Si está al borde, ¿hacia qué dirección?:', ['Sube', 'Baja'])

add_checkbox_options('Programa recomendado:', [
    'A1 nocturno (2h)',
    'A2 nocturno (2h)',
    'A2 4h intensivo mañana',
    'B1 (1h50m × 2 bloques)',
    'B2 (4h)',
    'Otro:  ____________________________________'
])

add_separator()

# ===== SECCIÓN 3 =====
add_heading('3. DOS FORTALEZAS CONCRETAS', level=2)
add_intro_quote(
    'Específico, no genérico. NO escribir "tiene buen vocabulario" — escribir '
    '"usa correctamente phrasal verbs como pick up y figure out".'
)

add_writing_block('Fortaleza 1:', lines=2)
add_quote_block('Cita textual o ejemplo escuchado (si lo hubo):', lines=2)

add_paragraph('')

add_writing_block('Fortaleza 2:', lines=2)
add_quote_block('Cita textual o ejemplo escuchado (si lo hubo):', lines=2)

add_separator()

# ===== SECCIÓN 4 =====
add_heading('4. UN PUNTO DE TRABAJO ESPECÍFICO', level=2)
add_intro_quote(
    'Lo que más se traba. NO escribir "le falta práctica" — escribir '
    '"confunde Past Simple con Present Perfect".'
)

add_writing_block('Punto de trabajo:', lines=2)
add_writing_block('¿Cómo lo va a resolver el programa? (1 línea):', lines=1)

add_separator()

# ===== SECCIÓN 5 =====
add_heading('5. MOTIVACIÓN PROFUNDA DETECTADA', level=2)

add_checkbox_options('¿Qué es lo que MÁS quiere lograr con el inglés?', [
    'Trabajo (ascenso / nuevo puesto / cambio de empresa)',
    'Viaje (vacaciones / migración / intercambio)',
    'Familia (hijos / pareja extranjera / nietos)',
    'Personal (autorrealización / superar trauma con el inglés)',
    'Negocio propio (clientes internacionales / marca personal)',
    'Estudios (maestría / certificación profesional)',
    'Otra:  ____________________________________'
])

add_quote_block('Detalle específico (qué dijo el prospecto que dejó claro su motivación):', lines=3)

add_separator()

# ===== SECCIÓN 6 =====
add_heading('6. LECTURA EMOCIONAL', level=2)
add_intro_quote('¿Cómo se siente el prospecto al cerrar la evaluación?')

add_checkbox_options('Estado emocional al cerrar:', [
    'CALIENTE — listo para cerrar. Mostró emoción, hizo preguntas sobre cómo arrancar, expresó urgencia.',
    'TIBIO — interesado pero con duda específica',
    'FRÍO — desinflado o desconfiado. Necesita 2do contacto antes de poder ofrecer programa.'
])

add_paragraph('')

add_checkbox_options('Si TIBIO, ¿cuál es la duda específica?', [
    'Costo / presupuesto',
    'Tiempo / agenda',
    'Método / si va a funcionar',
    'Decisión familiar (necesita consultar pareja/papás)',
    'Otra:  ____________________________________'
])

add_writing_block('¿Por qué tu lectura emocional? (1-2 frases):', lines=2)

add_separator()

# ===== SECCIÓN 7 =====
add_heading('7. COMPROMISO TEMPORAL HECHO AL PROSPECTO', level=2)
add_intro_quote(
    'Esta es la promesa que cerraste con el prospecto sobre cuándo lo contacta el asesor. '
    'Si el asesor no cumple esto, perdiste el cierre.'
)

add_checkbox_options('Le dije al prospecto que:', [
    'El asesor lo contacta HOY antes de las  ____ : ____',
    'El asesor lo contacta en las próximas  ____  horas',
    'El asesor lo contacta mañana antes de las  ____ : ____',
    '(Solo presencial) El asesor entró al salón inmediatamente después',
    'Otro:  ____________________________________'
])

add_checkbox_options('Canal de contacto acordado:', [
    'WhatsApp', 'Llamada telefónica', 'Cita presencial', 'Videollamada'
])

add_separator()

# ===== SECCIÓN 8 =====
add_heading('8. FRASES TEXTUALES IMPORTANTES (oro puro para el asesor)', level=2)

add_quote_block('Frase que indica INTENCIÓN DE COMPRA (si la dijo):', lines=2)
add_intro_quote(
    'Ejemplos: "necesito tener inglés en 6 meses", "ya intenté con otra academia y no me funcionó", '
    '"estoy listo para invertir en mí".'
)

add_paragraph('')

add_quote_block('Frase que indica OBJECIÓN o DUDA (si la dijo):', lines=2)
add_intro_quote(
    'Ejemplos: "el próximo mes lo veo", "tengo que hablar con mi esposa", "me parece caro", '
    '"no sé si voy a tener tiempo".'
)

add_separator()

# ===== SECCIÓN 9 =====
add_heading('9. BRIEFING RÁPIDO PARA EL ASESOR (2-3 líneas)', level=2)
add_intro_quote('Resumen ejecutivo. Lo primero que el asesor lee antes de contactar.')

add_writing_block('', lines=3)

add_separator()

# ===== SECCIÓN 10 =====
add_heading('10. ANCLAS PARA EL PITCH DEL ASESOR', level=2)
add_intro_quote(
    '2-3 conexiones específicas que el asesor puede usar para personalizar su contacto. '
    'NO genérico.'
)

add_writing_block('Ancla 1 (relacionada con su profesión / contexto):', lines=2)
add_intro_quote(
    'Ejemplo: "Trabaja en marketing → mencionarle que clientes en B2 acceden a '
    'roles de marketing en multinacionales y BPOs bilingües."'
)

add_paragraph('')

add_writing_block('Ancla 2 (relacionada con su motivación profunda):', lines=2)
add_intro_quote(
    'Ejemplo: "Quiere ir a Canadá en 18 meses → mostrarle que el programa B2 lo deja '
    'con nivel para entrevista IELTS."'
)

add_paragraph('')

add_writing_block('Ancla 3 (relacionada con sus fortalezas detectadas):', lines=2)
add_intro_quote(
    'Ejemplo: "Tiene buena pronunciación → resaltarle que va a destacar más rápido '
    'que el promedio en producción oral."'
)

add_separator()

# ===== SECCIÓN 11 =====
add_heading('11. RECOMENDACIÓN URGENCIA', level=2)

add_checkbox_options('¿Qué tan rápido debe contactar el asesor?', [
    'INMEDIATO (<30 min) — prospecto caliente, momentum alto',
    'HOY (mismo día) — interés real, pero no urgente',
    '24 HORAS (mañana) — tibio, dejar que procese',
    '48 HORAS+ (más adelante) — frío, necesita re-enganche con valor agregado'
])

add_writing_block('Si recomendaste contacto en >24h, ¿por qué?', lines=2)

add_separator()

# ===== SECCIÓN 12 =====
add_heading('12. NOTAS LIBRES', level=2)
add_intro_quote(
    'Cualquier observación que no entre en las secciones anteriores. Patrones detectados, '
    'alertas, contexto familiar/profesional importante, etc.'
)

add_writing_block('', lines=5)

add_separator()

# ===== CIERRE =====
add_heading('CIERRE DEL REPORTE', level=2)

add_field_table([
    ('Hora en que termino este reporte', ''),
    ('Hora en que envié este reporte al asesor', ''),
])

add_intro_quote(
    'Meta: llenar y enviar este reporte en menos de 5 min después de cerrar la evaluación. '
    'Más de 30 min = la información se diluye y pierde valor para el cierre.'
)

add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Firma evaluadora: ')
run.bold = True
p.add_run('_' * 60)

# Page break para instrucciones internas
doc.add_page_break()

# ===== INSTRUCCIONES PARA EL ASESOR =====
add_heading('INSTRUCCIONES PARA EL ASESOR (NO llenar la evaluadora)', level=2)

instructions_asesor = [
    '1. Léelo COMPLETO antes de contactar. No improvises.',
    '2. Usa las anclas (Sección 10) para personalizar tu primer mensaje. NO mensaje genérico.',
    '3. Cumple el compromiso temporal (Sección 7). Sin excusas.',
    '4. Si el prospecto es CALIENTE: llama, no escribas. Cierre se hace por voz, no por chat.',
    '5. Si es TIBIO con duda específica: trabaja esa duda directamente.',
    '6. Si es FRÍO: segundo contacto en 2-3 días con valor agregado (artículo, video, testimonio).',
]
for inst in instructions_asesor:
    p = doc.add_paragraph(inst)
    p.paragraph_format.left_indent = Cm(0.5)

add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Cierres reales = compromiso del asesor cumplido + información específica usada + decisión hoy.')
run.bold = True
run.italic = True

add_separator()

# ===== INSTRUCCIONES COORDINACIÓN =====
add_heading('INSTRUCCIONES PARA COORDINACIÓN (archivado)', level=2)

instructions_coord = [
    'Archivar TODOS los reportes (físicos o digitales) por mes.',
    'Cada lunes: revisar reportes de la semana anterior y verificar conversion rate (cuántos cerraron / cuántos se evaluaron).',
    'Si conversion <30%: revisar dónde se rompe (handoff tardío, mensaje genérico, no se trabajó objeción).',
    'Si un asesor tiene conversion <20% en 5+ evaluaciones: revisión 1-a-1 con la gerente comercial.',
]
for inst in instructions_coord:
    p = doc.add_paragraph(inst, style='List Bullet')

doc.save(DOCX_PATH)
print(f'OK: {DOCX_PATH}')
