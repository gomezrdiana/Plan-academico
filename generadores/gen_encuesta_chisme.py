#!/usr/bin/env python3
import os as _hos, sys as _hsys
_hsys.path.insert(0, _hos.path.dirname(_hos.path.dirname(_hos.path.abspath(__file__))))
"""Encuesta Conversacional Heiiu 2026 — modo chisme.
Genera Word (.docx) con 2 guiones cortos (Antiguos + Activos).
Foco: descubrir POR QUE se inscribieron + NPS al final."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(D, 'encuestas', 'Encuesta_Conversacional_Heiiu_2026.docx')

DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xA0, 0x20, 0x20)

doc = Document()

# ====== ESTILO BASE ======
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Margenes
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_title(text, size=20, color=DARK, align='left'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align == 'left' else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color


def add_subtitle(text, size=14, color=DARK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color


def add_section_header(text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = DARK


def add_para(text, size=11, italic=False, color=None, indent_left=0):
    p = doc.add_paragraph()
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_quote(text):
    """Caja con el script literal para la coordinadora."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'  "{text}"  ')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = DARK
    # Border izquierdo
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), '1A1A2E')
    pBdr.append(left)
    pPr.append(pBdr)


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'→ {text}')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_field(label):
    """Etiqueta + linea para escribir (4 lineas de espacio)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    r.font.bold = True
    r.font.size = Pt(11)
    p2 = doc.add_paragraph('_' * 110)
    p2.paragraph_format.space_after = Pt(0)
    p3 = doc.add_paragraph('_' * 110)
    p3.paragraph_format.space_after = Pt(0)
    p4 = doc.add_paragraph('_' * 110)
    p4.paragraph_format.space_after = Pt(0)
    p5 = doc.add_paragraph('_' * 110)
    p5.paragraph_format.space_after = Pt(6)


def add_field_small(label):
    """Etiqueta inline + un espacio corto (para datos rapidos)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{label}  ')
    r.font.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run('_' * 70)
    r2.font.size = Pt(11)


def add_long_field(label, lines=8):
    """Etiqueta + N lineas para respuesta literal larga."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    r.font.bold = True
    r.font.size = Pt(11)
    for _ in range(lines):
        pl = doc.add_paragraph('_' * 110)
        pl.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_nps_scale(prefix='NPS'):
    """Linea para marcar NPS 0-10."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'{prefix}:  ')
    r.font.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run('0   1   2   3   4   5   6   7   8   9   10')
    r2.font.size = Pt(13)
    r2.font.bold = True
    # Linea para circulo
    p2 = doc.add_paragraph()
    rp = p2.add_run('(marca con un circulo el numero)')
    rp.font.italic = True
    rp.font.size = Pt(9)
    rp.font.color.rgb = GRAY


def add_horizontal_rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '888888')
    pBdr.append(bot)
    pPr.append(pBdr)


def page_break():
    doc.add_page_break()


# ====================================================================
# PORTADA
# ====================================================================

add_title('HEIIU ENGLISH ACADEMY', size=18)
add_subtitle('Encuesta Conversacional · 2026')
add_para('Documento para la Coordinadora — Llamada en modo conversacion', italic=True, color=GRAY)

add_section_header('Proposito')
add_para(
    'Descubrir el verdadero motivo por el cual los estudiantes se inscribieron en Heiiu — '
    'tanto los que ya pasaron por el programa como los que estan activos hoy. '
    'La pregunta principal es "POR QUE TE INSCRIBISTE". Al final, una sola pregunta NPS.'
)

add_section_header('Modalidad — Modo "chisme"')
add_para('Esta NO es una encuesta formal. Es una conversacion casual. Reglas:')
add_para('· Tono conversacional, como si estuviera tomando un cafe con el estudiante.', indent_left=0.5)
add_para('· No leas las preguntas literal — adaptalas al flujo de la conversacion.', indent_left=0.5)
add_para('· Deja que el estudiante divague. Las historias son lo que estamos buscando.', indent_left=0.5)
add_para('· Anota frases textuales cuando puedas — son oro para el analisis.', indent_left=0.5)
add_para('· No defiendas Heiiu. No corrijas. Solo escucha y registra.', indent_left=0.5)

add_section_header('Alcance')
add_para('Total: 12 llamadas')
add_para('· 6 estudiantes ANTIGUOS (graduados o desertores) — usar Guion A', indent_left=0.5)
add_para('· 6 estudiantes ACTIVOS (matriculados hoy) — usar Guion B', indent_left=0.5)
add_para('Duracion esperada por llamada: 10-15 minutos.')

add_section_header('Confidencialidad')
add_para(
    'Al abrir cada llamada, recuerdale al estudiante que la conversacion es solo para revision '
    'interna de Heiiu y que su nombre no se compartira con su profesor sin su consentimiento.'
)

page_break()

# ====================================================================
# GUION A — ANTIGUOS
# ====================================================================

add_title('GUION A — Estudiantes Antiguos', size=16)
add_para('Para estudiantes ya graduados del programa o que se retiraron antes de terminar.', italic=True, color=GRAY)

add_section_header('Datos del estudiante (completar antes de llamar)')
add_field_small('Nombre completo:')
add_field_small('Telefono:')
add_field_small('Nivel alcanzado:')
add_field_small('Anio en que estudio:')
add_field_small('Graduado o desertor:')
add_field_small('Fecha de la llamada:')
add_field_small('Coordinadora:')

add_section_header('1. Apertura calida (1 min)')
add_quote(
    'Hola [nombre], te habla [coordinadora] de Heiiu. Como estas? Te llamo porque '
    'quiero hacerte un par de preguntas sobre tu experiencia, sin compromiso. '
    'Tienes 10 minutos para mi?'
)
add_quote(
    'Te aclaro: esto NO es una encuesta formal, es mas como una conversacion. '
    'Solo es para revision interna de Heiiu, tu nombre no se comparte. '
    'Puedes ser totalmente honesto.'
)
add_note('Si dice que no tiene tiempo: ofrecer reagendar. NO insistir.')

add_section_header('2. La pregunta central — el POR QUE')
add_quote(
    'Cuentame, regresemos a ese momento — que fue lo que te llevo a inscribirte en Heiiu en su momento? '
    'Que estabas buscando, que esperabas, que te animo a entrar?'
)
add_note('Esta es la pregunta mas importante. Dejalo hablar. No interrumpas. Anota literal.')
add_long_field('Respuesta literal del estudiante:', lines=10)

add_section_header('3. Que encontro vs lo que esperaba')
add_quote(
    'Cuando ya entraste y empezaste, lo que te encontraste se parecio a lo que esperabas? '
    'En que se parecio y en que fue distinto?'
)
add_long_field('Respuesta literal:', lines=7)

add_section_header('4. Que se llevo / que le falto')
add_quote(
    'De todo lo que viviste en Heiiu, que es lo que te llevaste? Y por otro lado, '
    'que te quedo pendiente, que te hubiera gustado vivir y no viviste?'
)
add_long_field('Respuesta literal:', lines=8)

add_horizontal_rule()
add_section_header('5. Cierre — NPS')
add_quote(
    'Para cerrar, una sola pregunta — del 1 al 10, que tan probable es que recomiendes Heiiu '
    'a un amigo o familiar? Donde 1 es no, ni de chiste, y 10 es lo recomendaria con los ojos cerrados.'
)
add_nps_scale('NPS final (1-10)')
add_long_field('Por que ese numero?', lines=4)

add_section_header('Cierre de la llamada')
add_quote(
    '[Nombre], muchas gracias por tu tiempo y tu honestidad. Lo que me contaste va a servir mucho. '
    'Si en algun momento te interesa volver a ver Heiiu, las puertas estan abiertas. Que estes bien.'
)
add_note('Si el estudiante quedo molesto o herido durante la conversacion: cierre con disculpa sincera, sin invitacion.')

add_section_header('Notas finales de la coordinadora')
add_field('Frase textual del estudiante que mejor resume su experiencia:')
add_field('Estado emocional al cierre de la llamada (calido / neutral / frio / molesto):')
add_field('Accion recomendada (escalar a directora si hay alerta, etc.):')

page_break()

# ====================================================================
# GUION B — ACTIVOS
# ====================================================================

add_title('GUION B — Estudiantes Activos', size=16)
add_para('Para estudiantes actualmente matriculados (2026).', italic=True, color=GRAY)

add_section_header('Datos del estudiante (completar antes de llamar)')
add_field_small('Nombre completo:')
add_field_small('Telefono:')
add_field_small('Nivel actual:')
add_field_small('Profesor:')
add_field_small('Semanas en la academia:')
add_field_small('Fecha de la llamada:')
add_field_small('Coordinadora:')

add_section_header('1. Apertura calida (1 min)')
add_quote(
    'Hola [nombre], te habla [coordinadora] de Heiiu. Como estas? Te llamo porque '
    'quiero conversar contigo unos minutos sobre tu experiencia hasta ahora. Tienes 10 minutos?'
)
add_quote(
    'Te aclaro: esto NO es una encuesta formal, es mas bien una conversacion. '
    'Es solo para revision interna, tu nombre no se va a compartir con tu profesor. '
    'Puedes ser totalmente honesto.'
)
add_note('No es para vender mas niveles. Si saca el tema, redirigir: "de eso hablamos despues".')

add_section_header('2. La pregunta central — el POR QUE')
add_quote(
    'Cuentame, que fue lo que te llevo a inscribirte en Heiiu? Que estabas buscando, que '
    'esperabas, por que decidiste empezar ahora y no antes?'
)
add_note('La pregunta mas importante. Dejalo hablar. No interrumpas. Anota literal.')
add_long_field('Respuesta literal del estudiante:', lines=10)

add_section_header('3. Como te esta pareciendo hasta ahora')
add_quote(
    'Llevas [X semanas] con nosotros. Como te esta pareciendo? Que es lo que sientes '
    'que esta funcionando, y que sientes que no?'
)
add_long_field('Respuesta literal:', lines=8)

add_section_header('4. Sorpresas — positivas y negativas')
add_quote(
    'Algo que te haya sorprendido en estas semanas? Algo que dijiste "wow, esto no me lo esperaba" — '
    'sea bueno o sea malo. Que ha sido?'
)
add_long_field('Respuesta literal:', lines=7)

add_horizontal_rule()
add_section_header('5. Cierre — NPS')
add_quote(
    'Para cerrar, una sola pregunta — del 1 al 10, hoy, en este momento, que tan probable es '
    'que recomiendes Heiiu a un amigo o familiar? Donde 1 es no, ni de chiste, y 10 es lo '
    'recomendaria con los ojos cerrados.'
)
add_nps_scale('NPS final (1-10)')
add_long_field('Por que ese numero?', lines=4)

add_section_header('Cierre de la llamada')
add_quote(
    '[Nombre], muchas gracias por la conversacion. Lo que me contaste me sirve mucho '
    'para entender como vamos. Cualquier cosa que quieras comentar despues, escribeme directo. '
    'Que sigas teniendo buena experiencia.'
)
add_note('Si el estudiante menciono algo critico (considera retirarse, mala experiencia, profesor que falla): notificar a la directora HOY mismo.')

add_section_header('Notas finales de la coordinadora')
add_field('Frase textual del estudiante que mejor resume su experiencia hoy:')
add_field('Estado emocional al cierre (motivado / neutral / con dudas / frustrado):')
add_field('Accion recomendada (escalar a directora si hay alerta roja):')

# ====== GUARDAR ======
doc.save(OUT)
print(f'OK: {OUT}')
